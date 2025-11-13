"""
Document processing service for RAG
Handles file upload, parsing, and text chunking
"""

import os
import io
import logging
from pathlib import Path
from typing import List, Tuple, Optional
import PyPDF2
from docx import Document as DocxDocument
import tiktoken

from app.models import Document, DocumentType, DocumentStatus, DocumentChunk
from app.database import Session

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing documents for RAG"""

    def __init__(self, upload_dir: str = "data/uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

        # Token counter for chunk sizing
        self.encoding = tiktoken.get_encoding("cl100k_base")

        # Chunk configuration
        self.chunk_size = 500  # tokens per chunk
        self.chunk_overlap = 50  # overlapping tokens between chunks

    async def save_uploaded_file(
        self,
        file_content: bytes,
        filename: str,
        db: Session
    ) -> Document:
        """
        Save uploaded file and create database record

        Args:
            file_content: Binary file content
            filename: Original filename
            db: Database session

        Returns:
            Document: Created document record
        """
        # Determine file type
        file_ext = Path(filename).suffix.lower().lstrip('.')

        try:
            doc_type = DocumentType(file_ext)
        except ValueError:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Generate unique file path
        file_path = self.upload_dir / f"{uuid.uuid4()}_{filename}"

        # Save file
        file_path.write_bytes(file_content)
        logger.info(f"Saved file to {file_path}")

        # Create document record
        document = Document(
            name=filename,
            file_path=str(file_path),
            doc_type=doc_type,
            status=DocumentStatus.PENDING,
            file_size=len(file_content)
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return document

    def extract_text(self, file_path: str, doc_type: DocumentType) -> str:
        """
        Extract text from document based on type

        Args:
            file_path: Path to file
            doc_type: Document type

        Returns:
            str: Extracted text
        """
        if doc_type == DocumentType.PDF:
            return self._extract_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            return self._extract_docx(file_path)
        elif doc_type in (DocumentType.TXT, DocumentType.MD):
            return self._extract_text(file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text_parts = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"Extracted {len(page_text)} chars from page {page_num + 1}")

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} total chars from PDF")
        return full_text

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} chars from DOCX")
        return full_text

    def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        logger.info(f"Extracted {len(text)} chars from text file")
        return text

    def chunk_text(self, text: str) -> List[Tuple[str, int]]:
        """
        Split text into chunks with token overlap

        Args:
            text: Full document text

        Returns:
            List[Tuple[str, int]]: List of (chunk_text, token_count) tuples
        """
        # Encode text to tokens
        tokens = self.encoding.encode(text)
        total_tokens = len(tokens)

        chunks = []
        start_idx = 0

        while start_idx < total_tokens:
            # Get chunk of tokens
            end_idx = min(start_idx + self.chunk_size, total_tokens)
            chunk_tokens = tokens[start_idx:end_idx]

            # Decode back to text
            chunk_text = self.encoding.decode(chunk_tokens)
            chunk_size = len(chunk_tokens)

            chunks.append((chunk_text, chunk_size))

            # Move to next chunk with overlap
            start_idx += self.chunk_size - self.chunk_overlap

        logger.info(f"Created {len(chunks)} chunks from {total_tokens} tokens")
        return chunks

    async def process_document(
        self,
        document: Document,
        db: Session
    ) -> None:
        """
        Process document: extract text, create chunks, update status

        Args:
            document: Document to process
            db: Database session
        """
        try:
            # Update status
            document.status = DocumentStatus.PROCESSING
            db.commit()

            # Extract text
            logger.info(f"Processing document: {document.name}")
            text = self.extract_text(document.file_path, document.doc_type)

            if not text.strip():
                raise ValueError("No text content extracted from document")

            # Create chunks
            chunks = self.chunk_text(text)

            # Save chunks to database
            for idx, (chunk_text, token_count) in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=idx,
                    content=chunk_text,
                    token_count=token_count,
                    meta_data={}
                )
                db.add(chunk)

            # Update document
            document.chunks_count = len(chunks)
            document.status = DocumentStatus.INDEXED
            db.commit()

            logger.info(f"✅ Successfully processed document {document.name}: {len(chunks)} chunks created")

        except Exception as e:
            logger.error(f"Failed to process document {document.name}: {e}")
            document.status = DocumentStatus.FAILED
            document.error_message = str(e)
            db.commit()
            raise


# Import uuid at the top
import uuid

# Create singleton instance
document_processor = DocumentProcessor()
